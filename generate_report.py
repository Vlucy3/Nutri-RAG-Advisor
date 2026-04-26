from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- PAGE MARGINS ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# --- STYLE HELPERS ---
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.runs[0].font.size = Pt(16)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.runs[0].font.size = Pt(13)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.runs[0].font.size = Pt(11)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def screenshot_box(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {label} ]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run.font.italic = True
    # Light grey shaded box via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(1)
    p.paragraph_format.right_indent = Inches(1)
    return p

def divider():
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

def inline_code(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return run

def blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

# ============================================================
# TITLE BLOCK
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Her-RAG Svetovalka 2026")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Retrieval-Augmented Generation Web Application")
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta.add_run("Author: Lucija Vovk   |   GitHub: github.com/Vlucy3/Nutri-RAG-Advisor   |   Live: nutrition-rag-app.onrender.com")
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ============================================================
# PAGE 1 — APPLICATION OVERVIEW
# ============================================================
heading1("Page 1: Application Overview")

heading2("Topic Description")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.add_run("I chose ").font.size = Pt(11)
r = p.add_run("women's hormonal nutrition")
r.bold = True; r.font.size = Pt(11)
p.add_run(
    " as the topic for this project. As a young woman, I find the connection between "
    "the menstrual cycle, food choices, and mood genuinely interesting, and I wanted to "
    "build something I would actually use. The app covers four hormonal phases of the "
    "menstrual cycle, science-backed recipes for specific symptoms (fatigue, anxiety, "
    "brain fog, bloating), and the latest 2024–2025 research on macronutrients, gut "
    "health, micronutrient synergy, and clinical nutrition."
).font.size = Pt(11)

heading2("What the App Does")
p = doc.add_paragraph()
p.add_run("Her-RAG Advisor 2026").bold = True
p.runs[-1].font.size = Pt(11)
p.add_run(" is a RAG web application with four pages:").font.size = Pt(11)

bullet("Home — Introduction to the app and its purpose.")
bullet("Hormonal Search — Users type a nutrition question and the app returns the most relevant document chunks from the knowledge base, with a match percentage and source label.")
bullet("Mood-Prep Kitchen — Users select how they feel from a symptom list and the app retrieves the best-matching recipe with ingredients, instructions, and a scientific explanation.")
bullet("Stats — Shows the full knowledge base breakdown and a live chunking strategy comparison.")

heading2("Screenshots")
screenshot_box("INSERT SCREENSHOT — Hormonal Search page with a search query and results")
doc.add_paragraph()
screenshot_box("INSERT SCREENSHOT — Mood-Prep Kitchen page with selected symptoms and recipe result")

heading2("Links")
bullet("GitHub repository: https://github.com/Vlucy3/Nutri-RAG-Advisor")
bullet("Live Render deployment: https://nutrition-rag-app.onrender.com")

page_break()

# ============================================================
# PAGE 2 — TECHNICAL DETAILS
# ============================================================
heading1("Page 2: Technical Details")

heading2("Documents")
p = doc.add_paragraph()
p.add_run("The knowledge base contains ").font.size = Pt(11)
r = p.add_run("30 total documents")
r.bold = True; r.font.size = Pt(11)
p.add_run(" across three source types:").font.size = Pt(11)

add_table(
    ["Source", "Count", "Description"],
    [
        ["Science modules", "5",  "Short hormonal science facts (infradian rhythm, cortisol, PCOS, endometriosis, estrogen metabolism)"],
        ["Recipes",         "10", "Mood-targeted recipes with ingredients, instructions, and biochemical mechanisms"],
        ["Research documents", "15", "Longer markdown files adapted from WHO guidelines, Endocrine Society publications, Nature Metabolism (2024), and nutrition science literature"],
    ],
    col_widths=[1.4, 0.6, 3.8]
)

body(
    "All content was written in English. Research documents cover: menstrual cycle phase "
    "nutrition, gut microbiome and SCFAs, macronutrient metabolism, hydration kinetics, "
    "micronutrient synergy, plant-based bioenergetics, clinical diabetes nutrition, sports "
    "ergogenic aids, dietary guidelines 2020-2025, and practical eating tips."
)

# --- CHUNKING ---
heading2("Chunking Strategy")
p = doc.add_paragraph()
r = p.add_run("Chosen strategy: ")
r.bold = True; r.font.size = Pt(11)
inline_code(p, "chunk_size=600, chunk_overlap=100")

body(
    "I chose chunk_size=600 because my documents contain multi-sentence scientific "
    "explanations that lose meaning when split too finely. A chunk of 600 characters is "
    "large enough to keep a full hormonal phase description (e.g. the entire Luteal phase "
    "paragraph) in a single chunk, so retrieval returns complete, context-rich answers "
    "rather than isolated fragments."
)
body(
    "The chunk_overlap=100 ensures that sentences spanning a chunk boundary are not lost — "
    "if a key recommendation appears at the end of one chunk, it also appears at the start "
    "of the next, preventing gaps in retrieved context."
)

heading3("Comparison with Strategy B: chunk_size=150, chunk_overlap=20")
body(
    "I tested a second strategy using micro-chunks (150/20), which is visible live in the "
    "Stats page of the app. The results were clearly worse:"
)

blockquote(
    'Example query: "luteal phase and magnesium"\n\n'
    'Strategy A (600/100): Returned the full Luteal phase paragraph — "Progesterone '
    'dominates, raising the basal metabolic rate by 100-300 calories... Focus on '
    'Magnesium-rich foods (pumpkin seeds, spinach) to reduce PMS-related water retention '
    'and anxiety..." — full context with actionable dietary recommendation.\n\n'
    'Strategy B (150/20): Returned only a recipe fragment — "Magnesium (Cacao) + B6 '
    '(Banana) = Melatonin production for deep recovery sleep." — correct keyword match, '
    'but no hormonal context or reasoning.'
)
body(
    "Conclusion: Strategy A (600/100) consistently returns richer, more informative chunks. "
    "Strategy B produces more precise keyword matching but sacrifices the surrounding "
    "scientific context that makes the answer useful."
)

# --- EMBEDDING MODEL ---
heading2("Embedding Model")
p = doc.add_paragraph()
r = p.add_run("Final model used: ")
r.bold = True; r.font.size = Pt(11)
inline_code(p, "intfloat/multilingual-e5-small")

body(
    "I chose a multilingual model because the app is in Slovenian, a language not supported "
    "by the default English-only all-MiniLM-L6-v2. The final model choice required two "
    "iterations due to a deployment memory constraint:"
)

add_table(
    ["Model", "Parameters", "Runtime RAM", "Slovenian", "Render free tier"],
    [
        ["all-MiniLM-L6-v2 (default)",                       "22M",  "~90 MB",  "No",  "✓ fits"],
        ["paraphrase-multilingual-MiniLM-L12-v2 (1st attempt)", "118M", "~470 MB", "Yes", "✗ OOM crash"],
        ["intfloat/multilingual-e5-small (final)",            "33M",  "~117 MB", "Yes", "✓ fits"],
    ],
    col_widths=[2.4, 0.9, 0.9, 0.8, 0.9]
)

heading3("Deployment Issue Encountered")
body(
    "After switching to paraphrase-multilingual-MiniLM-L12-v2, Render's free web service "
    "(512 MB RAM) crashed on startup with exit code 2. The model's 118M parameters require "
    "~470 MB of RAM at runtime, which — combined with PyTorch, ChromaDB, and Streamlit — "
    "exceeded the memory limit."
)
body(
    "Solution: I switched to intfloat/multilingual-e5-small, which supports 100+ languages "
    "including Slovenian, uses the same 384-dimensional embedding space, and fits "
    "comfortably within the 512 MB limit at ~117 MB runtime RAM. The model is loaded once "
    "at startup with @st.cache_resource and forced to CPU for Render compatibility."
)

# --- INTERESTING FINDINGS ---
heading2("Interesting Findings")

heading3("Off-topic queries behave gracefully")
body(
    'Searching for unrelated terms like "avto" (car) or "vreme" (weather) still returns '
    "results — ChromaDB always returns k results — but with very low match percentages "
    "(~5–15%). The match % label on each result card makes this visible to the user so "
    "they can judge result quality themselves."
)

heading3("Symptom combinations improve recipe matching")
body(
    'Selecting a single symptom like "Utrujenost" returns a good recipe, but combining '
    '"Utrujenost" + "Megla v glavi" shifts the result toward energy-and-focus recipes '
    "(like the Frittata or Poke Bowl), which are semantically closer to the combined "
    "query vector. The multilingual model handles compound Slovenian terms well."
)

heading3("Metadata filtering was essential for the Kitchen page")
body(
    'Before adding the filter={"source": "recipe"} metadata filter to the Mood-Prep '
    "Kitchen, the app would sometimes return a research document chunk (e.g. a paragraph "
    "about fatty acids) instead of a recipe card. The filter guarantees the Kitchen page "
    "always shows structured, cook-ready results."
)

page_break()

# ============================================================
# PAGE 3 — REFLECTIONS
# ============================================================
heading1("Page 3: Reflections & Extensions")

heading2("What I Learned")
body(
    "This project tied together everything from the course in a concrete, deployable "
    "product. The most valuable insight was understanding how chunk_size is not just a "
    "technical parameter but a semantic design decision — it determines how much context "
    "travels with each retrieved result. I also learned that multilingual embedding models "
    "require almost no extra code to integrate but unlock an entirely different audience "
    "for the application. Finally, the deployment challenge taught me that memory "
    "constraints on free cloud tiers are a real engineering concern, not just a "
    "theoretical one — model selection must account for the target runtime environment."
)

heading2("What I Would Improve With More Time")
bullet(
    "Add LLM response generation: The app currently returns raw document chunks. "
    "Connecting the Anthropic API would allow Claude to synthesize a conversational "
    "answer from the retrieved chunks — making it true RAG rather than retrieval-only."
)
bullet(
    "Add user feedback: A thumbs up/down button on each result would create a feedback "
    "loop to measure retrieval quality over time."
)
bullet(
    "Expand the recipe database: 10 recipes cover common symptoms but miss many "
    "combinations. Adding 20–30 more would improve the Kitchen's usefulness."
)
bullet(
    "Hybrid search: Combining semantic search (ChromaDB) with keyword search (BM25) "
    "would improve results for specific ingredient names or exact scientific terms."
)

# --- SAVE ---
out = "Her_RAG_Report_v2.docx"
doc.save(out)
print(f"Saved: {out}")
